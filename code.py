from docx import Document
from datetime import datetime
import locale
import calendar

mycal =calendar.Calendar()
year = 2021
month = 3
locale.setlocale(locale.LC_TIME, 'es-ES')
d = datetime.now()

months_values = mycal.itermonthdays(year,month)

CENTER = 1
LEFT = 0
RIGHT = 2
month = d.strftime("%B")
currentYear = datetime.now().year

weeks_count = len(calendar.monthcalendar(2021,3))
document = Document()
title_paragraph = document.add_paragraph('Plan de trabajo mensual', )
title_paragraph.alignment = CENTER
paragraph = document.add_paragraph('Mes: ' + str(month.capitalize()) + ' '+  ' Año: ' + str(currentYear)  )

table = document.add_table(rows=weeks_count * 2, cols=7)
table.style = 'Table Grid'


for n in range(7):
    if n == 0:
        table.cell(0, n).text = 'Lunes '
    if n == 1:
        table.cell(0, n).text = 'Martes '
    if n == 2:
        table.cell(0, n).text = 'Miercoles '
    if n == 3:
        table.cell(0, n).text = 'Jueves '
    if n == 4:
        table.cell(0, n).text = 'Viernes '
    if n == 5:
        table.cell(0, n).text = 'Sabado '
    if n == 6:
        table.cell(0, n).text = 'Domingo '
    


""" for m in months_values:
    table.cell(0, 0).text = 'Lunes ' + str(m)

 """

document.save('word.docx')


